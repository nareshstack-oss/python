from __future__ import annotations

from io import BytesIO
from textwrap import wrap

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from .models import ResumeOutput


def markdown_bytes(resume: ResumeOutput) -> bytes:
    return resume.resume_markdown.encode("utf-8")


def txt_bytes(resume: ResumeOutput) -> bytes:
    return resume.resume_markdown.encode("utf-8")


def docx_bytes(resume: ResumeOutput) -> bytes:
    document = Document()
    lines = [line.rstrip() for line in resume.resume_markdown.splitlines()]

    for line in lines:
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def pdf_bytes(resume: ResumeOutput) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    x = 50
    y = height - 50
    line_height = 16

    for raw_line in resume.resume_markdown.splitlines():
        line = raw_line.strip()
        if not line:
            y -= line_height
        else:
            font_name = "Helvetica"
            font_size = 11
            if line.startswith("# "):
                font_name = "Helvetica-Bold"
                font_size = 18
                line = line[2:].strip()
            elif line.startswith("## "):
                font_name = "Helvetica-Bold"
                font_size = 14
                line = line[3:].strip()
            elif line.startswith("- "):
                line = f"• {line[2:].strip()}"

            pdf.setFont(font_name, font_size)
            max_chars = 95 if font_size <= 11 else 70
            for wrapped in wrap(line, width=max_chars):
                if y < 60:
                    pdf.showPage()
                    y = height - 50
                    pdf.setFont(font_name, font_size)
                pdf.drawString(x, y, wrapped)
                y -= line_height
        if y < 60:
            pdf.showPage()
            y = height - 50

    pdf.save()
    return buffer.getvalue()
